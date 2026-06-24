"""Document ingestion utilities for loading and processing files efficiently."""

from collections.abc import AsyncIterator, Coroutine
from itertools import groupby
from pathlib import Path
from typing import Annotated, Callable, ClassVar, Self, TypeAlias, cast

from anyio import to_thread
import docx
from docx.oxml.ns import qn
from langchain_core.documents import Document
import openpyxl
from pydantic.fields import Field, computed_field
from pydantic.functional_validators import AfterValidator, model_validator
from pydantic.main import BaseModel
from pydantic.types import FilePath
from pypdf import PdfReader

from backend.core.config import IngestSettings, settings


def verify_file_integrity(path: Path) -> Path:
    """Prevent file spoofing."""
    try:
        with path.open("rb") as f:
            head: bytes = f.read(4)
    except OSError as err:
        raise ValueError(
            f"File validation failed: Unable to read file bytes. Error: {str(err)}"
        )

    suffix = path.suffix.lower()

    if suffix == ".pdf" and not head.startswith(b"%PDF"):
        raise ValueError(
            "File signature mismatch: File extension claims to be .pdf but headers do not match specification."
        )

    elif suffix in {".docx", ".xlsx"} and not head.startswith(b"PK"):
        raise ValueError(
            f"File signature mismatch: '{suffix}' container structure is corrupted or invalid."
        )

    return path


# Type Aliases
VerifiedFilePath: TypeAlias = Annotated[
    FilePath, AfterValidator(func=verify_file_integrity)
]
MetadataValue: TypeAlias = (
    str | int | float | bool | None | list[object] | dict[str, object]
)
StrictMetadata: TypeAlias = dict[str, MetadataValue]


class DocumentIngestor(BaseModel):
    """Asynchronously ingests documents from various file formats with schema protection."""

    allowed_extensions: ClassVar[set[str]] = {
        ".pdf",
        ".docx",
        ".md",
        ".xlsx",
    }
    file_path: Annotated[
        VerifiedFilePath, Field(description="File path of the Document to be ingested.")
    ]
    config: IngestSettings = settings.ingest

    @computed_field
    @property
    def extension(self) -> str:
        return self.file_path.suffix.lower()

    @model_validator(mode="after")
    def file_extention_in_allowed_extentions(self) -> Self:
        if self.extension not in self.allowed_extensions:
            raise ValueError(
                f"Unsupported file extension: {self.extension}. Supported formats: {', '.join(sorted(self.allowed_extensions))}"
            )
        return self

    async def _load_pdf(self, source: str) -> list[Document]:
        extract_images = self.config.pdf_extract_images

        def _parse() -> list[Document]:
            reader = PdfReader(stream=self.file_path)
            total_pages = len(reader.pages)
            docs: list[Document] = []
            for i, page in enumerate(reader.pages):
                parts = [page.extract_text() or ""]
                if extract_images:
                    for img in page.images:
                        if img.data:
                            parts.append(f"[image:{img.name}]")
                page_content = "\n".join(filter(None, parts)).strip()
                if not page_content:
                    continue
                docs.append(
                    Document(
                        page_content=page_content,
                        metadata={
                            "source": source,
                            "file_type": "pdf",
                            "page": i + 1,
                            "total_pages": total_pages,
                        },
                    )
                )
            return docs

        return await to_thread.run_sync(_parse)

    async def _load_docx(self, source: str) -> list[Document]:
        def _parse() -> list[Document]:
            doc = docx.Document(docx=str(self.file_path))
            core_props = doc.core_properties
            base_meta = {
                "source": source,
                "file_type": "docx",
                "author": core_props.author or "",
                "title": core_props.title or "",
            }

            section_idx = 0
            para_sections: list[tuple[int, str]] = []  # (section_idx, text)
            for para in doc.paragraphs:
                sect_pr = para._p.find(  # pyright: ignore[reportUnknownMemberType, reportPrivateUsage, reportUnknownVariableType]
                    qn(tag="w:pPr")
                )
                if (
                    sect_pr is not None
                    and sect_pr.find(  # pyright: ignore[reportUnknownMemberType]
                        qn(tag="w:sectPr")
                    )
                    is not None
                ):
                    if para.text:
                        para_sections.append((section_idx, para.text))
                    section_idx += 1
                elif para.text:
                    para_sections.append((section_idx, para.text))

            if section_idx == 0:
                batch_size = 50
                all_texts = [t for _, t in para_sections]
                batches: list[list[str]] = [
                    all_texts[i : i + batch_size]
                    for i in range(0, max(len(all_texts), 1), batch_size)
                ]
                docs: list[Document] = [
                    Document(
                        page_content="\n".join(batch),
                        metadata={**base_meta, "page": page_num + 1},
                    )
                    for page_num, batch in enumerate(batches)
                    if batch
                ]
            else:
                docs = []
                for sec, group in groupby(para_sections, key=lambda x: x[0]):
                    texts = [t for _, t in group]
                    if texts:
                        docs.append(
                            Document(
                                page_content="\n".join(texts),
                                metadata={**base_meta, "page": sec + 1},
                            )
                        )

            if self.config.docx_include_tables:
                for table in doc.tables:
                    for row in table.rows:
                        row_text = "\t".join(c.text for c in row.cells if c.text)
                        if row_text and docs:
                            docs[-1].page_content += "\n" + row_text

            if self.config.docx_include_headers_footers:
                for section in doc.sections:
                    for hf in (section.header, section.footer):
                        for p in hf.paragraphs:
                            if p.text and docs:
                                docs[-1].page_content += "\n" + p.text

            return (
                docs
                if docs
                else [Document(page_content="", metadata={**base_meta, "page": 1})]
            )

        return await to_thread.run_sync(_parse)

    async def _load_md(self, source: str) -> list[Document]:
        text = await to_thread.run_sync(lambda: self.file_path.read_text("utf-8"))
        return [
            Document(
                page_content=text,
                metadata={
                    "source": source,
                    "file_type": "markdown",
                    "page": 1,
                },
            )
        ]

    async def _load_xlsx(self, source: str) -> list[Document]:
        include_empty_rows = self.config.xlsx_include_empty_rows

        def _parse() -> list[Document]:
            wb = openpyxl.load_workbook(
                filename=self.file_path, read_only=True, data_only=True
            )
            total_sheets = len(wb.sheetnames)
            docs: list[Document] = []
            for sheet_index, sheet in enumerate(wb.worksheets):
                rows: list[str] = []
                for row in sheet.iter_rows():
                    row_text = "\t".join(str(cell.value or "") for cell in row)
                    if include_empty_rows or row_text.strip():
                        rows.append(row_text)
                docs.append(
                    Document(
                        page_content="\n".join(rows),
                        metadata={
                            "source": source,
                            "file_type": "xlsx",
                            "page": sheet_index + 1,
                            "sheet": sheet.title,
                            "total_sheets": total_sheets,
                        },
                    )
                )
            return docs

        return await to_thread.run_sync(_parse)

    async def ingest_lazy(self) -> AsyncIterator[Document]:
        loaders: dict[str, Callable[[str], Coroutine[None, None, list[Document]]]] = {
            ".pdf": self._load_pdf,
            ".docx": self._load_docx,
            ".md": self._load_md,
            ".xlsx": self._load_xlsx,
        }
        documents = await loaders[self.extension](str(self.file_path))

        current_page_tracking = 1

        for document in documents:
            metadata: StrictMetadata = cast(StrictMetadata, document.metadata)

            resolved_page = self._extract_page_number(
                metadata, fallback_index=current_page_tracking
            )
            current_page_tracking = resolved_page

            metadata["filename"] = self.file_path.name
            metadata["page_number"] = resolved_page

            yield document

    async def ingest_async(self) -> list[Document]:
        return [doc async for doc in self.ingest_lazy()]

    @staticmethod
    def _extract_page_number(
        metadata: StrictMetadata,
        fallback_index: int,
    ) -> int:
        page_value: MetadataValue = metadata.get("page") or metadata.get("page_number")
        if page_value is None:
            return fallback_index

        try:
            val = int(str(page_value))
            return val if val > 0 else fallback_index
        except (ValueError, TypeError):
            return fallback_index
