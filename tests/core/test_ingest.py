from io import BytesIO
from pathlib import Path
from unittest.mock import AsyncMock

import docx as python_docx
import openpyxl
from langchain_core.documents import Document
import pytest
from pypdf import PdfWriter
from pytest_mock.plugin import MockerFixture

from backend.core.config import IngestSettings
from backend.core.ingest import DocumentIngestor, StrictMetadata


def _make_pdf(tmp_path: Path, text: str = "Hello PDF") -> Path:
    writer = PdfWriter()
    page = writer.add_blank_page(width=200, height=200)
    page.merge_page(page)
    buf = BytesIO()
    writer.write(buf)
    path = tmp_path / "test.pdf"
    path.write_bytes(buf.getvalue())
    return path


def _make_docx(
    tmp_path: Path,
    paragraphs: list[str] | None = None,
    table_rows: list[list[str]] | None = None,
    header_text: str = "",
) -> Path:
    doc = python_docx.Document()
    for text in paragraphs or ["Hello DOCX"]:
        doc.add_paragraph(text)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, val in enumerate(row):
                table.cell(r, c).text = val
    if header_text:
        doc.sections[0].header.paragraphs[0].text = header_text
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path


def _make_xlsx(
    tmp_path: Path, sheets: dict[str, list[list[str]]] | None = None
) -> Path:
    wb = openpyxl.Workbook()
    default_sheet = wb.active
    assert default_sheet is not None
    default_data = sheets or {"Sheet": [["A", "B"], ["1", "2"]]}
    first = True
    for name, rows in default_data.items():
        if first:
            default_sheet.title = name
            ws = default_sheet
            first = False
        else:
            ws = wb.create_sheet(name)
        for row in rows:
            ws.append(row)
    path = tmp_path / "test.xlsx"
    wb.save(str(path))
    return path


class TestDocumentIngestor:
    async def test_unsupported_extension(self, tmp_path: Path) -> None:
        invalid_file = tmp_path / "sample.txt"
        _ = invalid_file.write_text("data")

        with pytest.raises(ValueError) as exc:
            _ = DocumentIngestor(file_path=invalid_file)

        assert "Unsupported file extension" in str(exc.value)

    async def test_ingest_async_success(
        self,
        tmp_path: Path,
        mocker: MockerFixture,
    ) -> None:
        mock_docs = [
            Document(page_content="test content 1", metadata={"page": "1"}),
            Document(page_content="test content 2", metadata={"page": "2"}),
        ]

        pdf_file = tmp_path / "test.pdf"
        _ = pdf_file.write_bytes(data=b"%PDF-1.4\n%test")

        mock_load: AsyncMock = mocker.patch.object(
            target=DocumentIngestor,
            attribute="_load_pdf",
            return_value=mock_docs,
        )

        ingest: DocumentIngestor = DocumentIngestor(file_path=pdf_file)

        result: list[Document] = await ingest.ingest_async()

        assert result == mock_docs
        assert len(result) == 2
        mock_load.assert_called_once()

    async def test_ingest_lazy_enriches_metadata(
        self, tmp_path: Path, mocker: MockerFixture
    ) -> None:
        mock_docs = [Document(page_content="text", metadata={"page": 2})]
        path = tmp_path / "test.pdf"
        path.write_bytes(b"%PDF-1.4\n%test")

        mocker.patch.object(DocumentIngestor, "_load_pdf", return_value=mock_docs)
        ingestor = DocumentIngestor(file_path=path)

        result = await ingestor.ingest_async()

        assert result[0].metadata["filename"] == "test.pdf"
        assert result[0].metadata["page_number"] == 2


class TestLoadPdf:
    async def test_returns_one_doc_per_page(self, tmp_path: Path) -> None:
        path = _make_pdf(tmp_path)
        ingestor = DocumentIngestor(file_path=path)
        docs = await ingestor._load_pdf(
            source=str(path)
        )  # pyright: ignore[reportPrivateUsage]
        assert len(docs) == 1
        assert docs[0].metadata["file_type"] == "pdf"
        assert docs[0].metadata["total_pages"] == 1
        assert docs[0].metadata["page"] == 1

    async def test_extract_images_flag(self, tmp_path: Path) -> None:
        path = _make_pdf(tmp_path)
        config = IngestSettings(pdf_extract_images=True)
        ingestor = DocumentIngestor(file_path=path, config=config)
        docs = await ingestor._load_pdf(
            source=str(path)
        )  # pyright: ignore[reportPrivateUsage]
        assert len(docs) >= 1


class TestLoadDocx:
    async def test_basic_paragraphs(self, tmp_path: Path) -> None:
        path = _make_docx(tmp_path, paragraphs=["Hello", "World"])
        ingestor = DocumentIngestor(file_path=path)
        docs = await ingestor._load_docx(
            source=str(path)
        )  # pyright: ignore[reportPrivateUsage]
        assert len(docs) == 1
        assert "Hello" in docs[0].page_content
        assert "World" in docs[0].page_content
        assert docs[0].metadata["file_type"] == "docx"
        assert docs[0].metadata["page"] == 1

    async def test_includes_tables_when_enabled(self, tmp_path: Path) -> None:
        path = _make_docx(tmp_path, table_rows=[["Col1", "Col2"], ["Val1", "Val2"]])
        config = IngestSettings(docx_include_tables=True)
        ingestor = DocumentIngestor(file_path=path, config=config)
        docs = await ingestor._load_docx(
            source=str(path)
        )  # pyright: ignore[reportPrivateUsage]
        assert "Col1" in docs[0].page_content

    async def test_excludes_tables_when_disabled(self, tmp_path: Path) -> None:
        path = _make_docx(tmp_path, table_rows=[["Secret", "Data"]])
        config = IngestSettings(docx_include_tables=False)
        ingestor = DocumentIngestor(file_path=path, config=config)
        docs = await ingestor._load_docx(
            source=str(path)
        )  # pyright: ignore[reportPrivateUsage]
        assert "Secret" not in docs[0].page_content

    async def test_includes_headers_footers_when_enabled(self, tmp_path: Path) -> None:
        path = _make_docx(tmp_path, header_text="MY HEADER")
        config = IngestSettings(docx_include_headers_footers=True)
        ingestor = DocumentIngestor(file_path=path, config=config)
        docs = await ingestor._load_docx(
            source=str(path)
        )  # pyright: ignore[reportPrivateUsage]
        assert "MY HEADER" in docs[0].page_content

    async def test_excludes_headers_footers_when_disabled(self, tmp_path: Path) -> None:
        path = _make_docx(tmp_path, header_text="MY HEADER")
        config = IngestSettings(docx_include_headers_footers=False)
        ingestor = DocumentIngestor(file_path=path, config=config)
        docs = await ingestor._load_docx(
            source=str(path)
        )  # pyright: ignore[reportPrivateUsage]
        assert "MY HEADER" not in docs[0].page_content


class TestLoadMd:
    async def test_returns_single_doc(self, tmp_path: Path) -> None:
        path = tmp_path / "test.md"
        path.write_text("# Hello\nWorld")
        ingestor = DocumentIngestor(file_path=path)
        docs = await ingestor._load_md(
            source=str(path)
        )  # pyright: ignore[reportPrivateUsage]
        assert len(docs) == 1
        assert "# Hello" in docs[0].page_content
        assert docs[0].metadata["file_type"] == "markdown"
        assert docs[0].metadata["page"] == 1


class TestLoadXlsx:
    async def test_returns_one_doc_per_sheet(self, tmp_path: Path) -> None:
        path = _make_xlsx(tmp_path, sheets={"A": [["1"]], "B": [["2"]]})
        ingestor = DocumentIngestor(file_path=path)
        docs = await ingestor._load_xlsx(
            source=str(path)
        )  # pyright: ignore[reportPrivateUsage]
        assert len(docs) == 2
        assert docs[0].metadata["sheet"] == "A"
        assert docs[1].metadata["sheet"] == "B"
        assert docs[0].metadata["total_sheets"] == 2
        assert docs[0].metadata["file_type"] == "xlsx"

    async def test_excludes_empty_rows_by_default(self, tmp_path: Path) -> None:
        path = _make_xlsx(tmp_path, sheets={"Sheet": [["data"], [], ["more"]]})
        ingestor = DocumentIngestor(file_path=path)
        docs = await ingestor._load_xlsx(
            source=str(path)
        )  # pyright: ignore[reportPrivateUsage]
        lines = [l for l in docs[0].page_content.splitlines() if l.strip()]
        assert len(lines) == 2

    async def test_includes_empty_rows_when_enabled(self, tmp_path: Path) -> None:
        path = _make_xlsx(tmp_path, sheets={"Sheet": [["data"], [], ["more"]]})
        config = IngestSettings(xlsx_include_empty_rows=True)
        ingestor = DocumentIngestor(file_path=path, config=config)
        docs = await ingestor._load_xlsx(
            source=str(path)
        )  # pyright: ignore[reportPrivateUsage]
        assert len(docs[0].page_content.splitlines()) == 3


class TestPageNumberExtraction:
    @pytest.mark.parametrize(
        argnames="metadata_dict,fallback_idx,expected_page",
        argvalues=[
            ({"page": "3"}, 1, 3),
            ({"page_number": 5}, 2, 5),
            ({}, 4, 4),
        ],
    )
    def test_extracts_page_number(
        self, metadata_dict: StrictMetadata, fallback_idx: int, expected_page: int
    ) -> None:
        assert (
            DocumentIngestor._extract_page_number(  # pyright: ignore[reportPrivateUsage]
                metadata=metadata_dict, fallback_index=fallback_idx
            )
            == expected_page
        )
